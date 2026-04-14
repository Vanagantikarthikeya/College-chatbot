
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3C, 0x6E)   # headings
TEAL   = RGBColor(0x00, 0x7B, 0x83)   # sub-headings
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF0, 0xF4, 0xF8)   # table header bg
MGRAY  = RGBColor(0x60, 0x60, 0x60)   # caption / label

# ── Helper: shade a table cell ────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Generic paragraph helpers ─────────────────────────────────────
def add_paragraph(text='', bold=False, italic=False, size=11,
                  color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color
    return p

def add_bullet(text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(18 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    return p

def add_code_block(text):
    """Light grey shaded paragraph for code / architecture diagrams."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EFF3F7')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return p

# ── Section heading helpers ───────────────────────────────────────
def add_main_heading(number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{number}.  {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = WHITE
    # navy shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1A3C6E')
    pPr.append(shd)
    return p

def add_sub_heading(title, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = TEAL
    return p

def add_label(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = BLACK
    return p

# ── Generic table builder ─────────────────────────────────────────
def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, '1A3C6E')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    # data rows
    for ridx, row_data in enumerate(rows):
        row = t.add_row()
        fill = 'F0F4F8' if ridx % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            shade_cell(cell, fill)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()   # spacing after table
    return t

# ═════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("AI-POWERED COLLEGE HELPDESK\nCHATBOT SYSTEM")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

add_paragraph("Pallavi Engineering College, Hyderabad", bold=True, size=13,
              color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
add_paragraph("Affiliated to JNTUH  |  Approved by AICTE", size=11,
              color=MGRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("Academic Year: 2025 – 2026", italic=True, size=11,
              color=MGRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
add_paragraph("Comprehensive Project Documentation", bold=True, size=12,
              color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  1. PROJECT BASIC DETAILS
# ═════════════════════════════════════════════════════════════════
add_main_heading(1, "PROJECT BASIC DETAILS")

add_label("Project Title", "AI-Powered College Helpdesk Chatbot System")
add_label("Domain", "Artificial Intelligence / Natural Language Processing / Web Application")
add_label("Institution", "Pallavi Engineering College, Hyderabad")
add_label("Affiliated University", "Jawaharlal Nehru Technological University, Hyderabad (JNTUH)")
add_label("Academic Year", "2025 – 2026")

add_sub_heading("Problem Statement")
add_paragraph(
    "Students at engineering colleges frequently struggle to obtain timely and accurate information "
    "regarding admissions, examinations, fee structures, library services, placements, and other "
    "institutional services. Traditional manual helpdesks operate only during office hours and are "
    "inconsistent, resource-intensive, and unable to scale with increasing student populations. This "
    "creates friction, delays, and dissatisfaction among students who require immediate answers."
)

add_sub_heading("Objective of the Project")
for obj in [
    "Develop an intelligent, web-based chatbot capable of answering college-specific student queries in natural language.",
    "Train an MLP neural network on institution-specific FAQ patterns for high-accuracy intent classification.",
    "Implement a semantic fallback using Sentence Transformers for paraphrased or novel queries.",
    "Provide a dynamic document search layer over admin-uploaded institutional documents (PDF, DOCX, TXT).",
    "Build an admin panel for managing FAQs, uploading documents, and monitoring system performance analytics.",
    "Log all unanswered queries to continuously identify and close knowledge gaps.",
]:
    add_bullet(obj)

add_sub_heading("Motivation Behind the Project")
for m in [
    "24/7 Availability: Students need answers outside of regular office hours.",
    "Reduced Admin Burden: Automate repetitive query handling to free staff for complex tasks.",
    "Consistent Information: Eliminate variability in responses across different staff members.",
    "Continuous Improvement: Unanswered query logs create a data-driven feedback loop.",
    "Digital Transformation: Modernise the college helpdesk to align with student expectations.",
    "AI in Education: Demonstrate a practical real-world application of machine learning.",
]:
    add_bullet(m)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  2. ABSTRACT
# ═════════════════════════════════════════════════════════════════
add_main_heading(2, "ABSTRACT")
add_paragraph(
    "The AI-Powered College Helpdesk Chatbot is an intelligent web application designed and developed "
    "for Pallavi Engineering College to facilitate automated, round-the-clock student assistance. The "
    "system employs a two-layer hybrid architecture that combines a trained Multi-Layer Perceptron (MLP) "
    "neural network for intent classification with a semantic sentence embedding fallback using the "
    "all-MiniLM-L6-v2 Sentence Transformer model. This dual strategy enables the chatbot to handle a "
    "wide variety of natural language queries with both high accuracy and semantic awareness, even when "
    "students use different vocabulary or phrasing from the training data."
    "\n\n"
    "When neither the static knowledge base (FAQ intents) nor the embedding similarity search produces a "
    "confident answer, the system automatically transitions to a dynamic document search layer. This layer "
    "queries admin-uploaded institutional documents through chunked keyword retrieval, returning the most "
    "relevant excerpt along with proper source attribution. All queries that cannot be resolved are logged "
    "to a structured JSON file, enabling administrators to identify knowledge gaps and improve coverage "
    "systematically over time."
    "\n\n"
    "The administrative backend, protected by session-based authentication, provides a comprehensive "
    "dashboard for managing FAQ intents, uploading and toggling institutional documents, and viewing "
    "live performance analytics including success rates, answer source distribution, and frequent "
    "unanswered queries. Built entirely using Flask (Python), Jinja2 HTML templating, scikit-learn, "
    "NLTK, and a lightweight JSON file store, the system is portable, easy to deploy, and requires no "
    "proprietary database infrastructure. This project demonstrates the effective integration of Natural "
    "Language Processing, machine learning, document intelligence, and web development to solve a "
    "tangible institutional challenge."
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  3. INTRODUCTION
# ═════════════════════════════════════════════════════════════════
add_main_heading(3, "INTRODUCTION")

add_sub_heading("Background of the Problem")
add_paragraph(
    "Engineering colleges manage thousands of enrolled students who continuously generate diverse queries "
    "spanning admissions, academic schedules, examination fees, library services, placement drives, "
    "hostel facilities, and official contact information. This information is often scattered across "
    "physical notice boards, static college websites, and administrative offices — leading to an "
    "inconsistent and time-consuming information retrieval experience for students."
)
add_paragraph(
    "During peak periods such as exam registrations, admission seasons, and placement drives, "
    "administrative offices become overwhelmed with repetitive queries that could easily be automated. "
    "Meanwhile, students who cannot reach the helpdesk in time face academic risks from missing "
    "deadlines and misinformation."
)

add_sub_heading("Importance of the Project")
for imp in [
    "Provides instant, consistent responses across all student queries — 24 hours a day, 7 days a week.",
    "Scales effortlessly — a single chatbot can serve thousands of concurrent students.",
    "Creates an intelligent feedback loop through unanswered query logging, driving continuous improvement.",
    "Reduces cognitive and operational load on administrative and academic staff.",
    "Allows non-technical admins to update the knowledge base without any programming knowledge.",
    "Serves as a living demonstration of AI applied to real-world educational infrastructure.",
]:
    add_bullet(imp)

add_sub_heading("Overview of the Solution")
add_paragraph(
    "The proposed system is a Python Flask-based web application that integrates three layers of "
    "intelligence into a single chatbot interface:"
)
for item in [
    "Layer 1a — MLP Neural Network: A trained intent classifier that maps user input to one of 40+ predefined FAQ categories with confidence scoring.",
    "Layer 1b — Sentence Embeddings: A semantic similarity fallback using the all-MiniLM-L6-v2 transformer model for paraphrase-resilient matching.",
    "Layer 2 — Document Search: A keyword overlap search over admin-uploaded institutional documents (PDF, DOCX, TXT) chunked into retrievable segments.",
    "Admin Dashboard: A secure panel for FAQ management, document ingestion, and performance analytics.",
]:
    add_bullet(item)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  4. EXISTING SYSTEM
# ═════════════════════════════════════════════════════════════════
add_main_heading(4, "EXISTING SYSTEM")

add_sub_heading("Current Systems Available")
add_table(
    ["System Type", "Description"],
    [
        ["Manual Help Desks", "Staff-operated inquiry counters in administrative offices during working hours only"],
        ["College Websites", "Static HTML pages with FAQ sections, notice boards, and downloadable PDFs"],
        ["Rule-Based Chatbots", "Simple keyword-matching bots with pre-programmed exact-match responses"],
        ["Generic AI Assistants", "General tools like ChatGPT with no institution-specific knowledge tuning"],
        ["Email / Phone Support", "Asynchronous communication with long response times and no central log"],
    ],
    col_widths=[2.0, 4.5]
)

add_sub_heading("Limitations and Drawbacks")
for lim in [
    "Not available 24/7 — students cannot get answers outside office working hours.",
    "Inconsistent responses — different staff members may provide conflicting information.",
    "No natural language understanding — keyword bots fail when students rephrase questions.",
    "Information silos — important documents are not searchable or retrievable in real-time.",
    "No analytics — no mechanism to track what students are asking or system failures.",
    "Manual updates — updating FAQ content requires technical knowledge or website access.",
    "No feedback loop — there is no system to identify what information is most needed.",
    "High cost at scale — serving thousands of queries requires proportional staff increase.",
]:
    add_bullet(lim)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  5. PROPOSED SYSTEM
# ═════════════════════════════════════════════════════════════════
add_main_heading(5, "PROPOSED SYSTEM")

add_sub_heading("What is the New System")
add_paragraph(
    "The proposed system is an AI-powered, two-layer hybrid chatbot web application tailored specifically "
    "for college environments. It combines a trained neural network (MLP) for intent classification, "
    "semantic sentence embeddings for paraphrase handling, and a dynamic document retrieval layer for "
    "admin-uploaded institutional content. All unanswered queries are logged and surfaced to admins "
    "through an analytics dashboard."
)

add_sub_heading("Proposed vs. Existing System Comparison")
add_table(
    ["Feature", "Existing System", "Proposed System"],
    [
        ["Availability", "Office hours only", "24 / 7"],
        ["Response Consistency", "Variable (staff-dependent)", "Uniform (AI-generated)"],
        ["Natural Language Processing", "None", "MLP + Sentence Embeddings"],
        ["Document Search", "Manual lookup", "Automated keyword chunk retrieval"],
        ["Admin Analytics", "None", "Live dashboard with success rate"],
        ["Knowledge Updates", "Technical / manual", "Admin-friendly web panel"],
        ["Unanswered Query Tracking", "None", "JSON logging with frequency analysis"],
        ["Scalability", "Limited by staff count", "Handles thousands simultaneously"],
    ],
    col_widths=[2.2, 2.0, 2.4]
)

add_sub_heading("Key Innovations")
for item in [
    "Two-Layer Hybrid Architecture: MLP classifier + document fallback ensures near-complete coverage.",
    "Semantic Embedding Fallback: Sentence Transformer cosine similarity handles paraphrased or novel queries.",
    "Admin-Driven Dynamic Knowledge: Admins upload PDFs and manage FAQs without any coding.",
    "Intelligent Query Logging: System surfaces frequent failures to admins for targeted improvement.",
    "Synonym Normalization: Pre-processing layer standardises synonyms (e.g., 'register' → 'apply') before ML inference.",
    "Scanned PDF Detection: Warns admins if an uploaded PDF contains no extractable text.",
]:
    add_bullet(item)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  6. SYSTEM ARCHITECTURE
# ═════════════════════════════════════════════════════════════════
add_main_heading(6, "SYSTEM ARCHITECTURE")

add_sub_heading("Detailed Explanation")
add_paragraph(
    "The system follows a layered, request-response architecture hosted on a Flask web server. "
    "The architecture comprises three main tiers: a Presentation Tier (HTML/CSS/JS frontend), "
    "an Application Tier (Flask routing + chatbot engine), and a Data Tier (JSON flat-file store)."
)

add_sub_heading("Architecture Diagram (Text Representation)")
add_code_block(
"""Student Browser
    │
    ▼  HTTP GET
Flask Web Server  (app.py)
    ├── Route: /           → main.html       (College Homepage)
    ├── Route: /helpdesk   → chatbot.html    (Student Chat UI)
    ├── Route: /get-response  [POST JSON]
    │       │
    │       └── chatbot_engine.chatbot_response()
    │               │
    │               ├── LAYER 1A: MLP Neural Network
    │               │       nlp_preprocessing  →  bag_of_words
    │               │       model.pkl          →  predict_proba
    │               │       Confidence ≥ 0.65  →  Return FAQ response
    │               │
    │               ├── LAYER 1B: Sentence Transformer Fallback
    │               │       encode(query)  →  cosine_similarity
    │               │       Score ≥ 0.50   →  Return embedded response
    │               │
    │               └── LAYER 2: Admin Document Search
    │                       document_ingestion.search_documents()
    │                       admin_static_knowledge.json  (chunk search)
    │                       Score ≥ 2 keywords  →  Return document response
    │                       ELSE: log query  +  return fallback message
    │
    └── Admin Routes  (/admin/*)
            ├── /admin/login       → Session authentication
            ├── /admin/dashboard   → Overview panel
            ├── /admin/knowledge   → FAQ + Document management
            ├── /admin/analytics   → Performance metrics
            └── /admin/logout      → Session clear"""
)

add_sub_heading("Components Involved")
add_table(
    ["Component", "File", "Role"],
    [
        ["Web Server & Routing", "app.py", "Flask routes, session management, admin logic"],
        ["Chatbot Engine", "chatbot_engine.py", "Two-layer hybrid NLP response engine"],
        ["NLP Preprocessor", "nlp_preprocessing.py", "Tokenisation and lemmatisation via NLTK"],
        ["Document Processor", "document_ingestion.py", "File parsing, chunking, keyword search"],
        ["Model Trainer", "train_model.py", "One-time offline MLP training script"],
        ["Intent Database", "intents.json", "FAQ patterns and responses (40+ intents)"],
        ["Trained Model", "model.pkl", "Serialised scikit-learn MLPClassifier"],
        ["Vocabulary", "words.pkl", "Lemmatised word list for Bag-of-Words"],
        ["Class Labels", "classes.pkl", "Intent tag list (output classes)"],
        ["Document Index", "data/admin_static_metadata.json", "Document metadata and active status"],
        ["Document Content", "data/admin_static_knowledge.json", "Chunked text for retrieval"],
        ["Query Log", "data/unanswered_queries.json", "Failed/partial response audit trail"],
    ],
    col_widths=[1.8, 2.3, 2.5]
)

add_sub_heading("Data Flow Explanation")
for step in [
    "Student submits query via chat UI → AJAX POST to /get-response with JSON body.",
    "Flask receives request → calls chatbot_engine.chatbot_response(message).",
    "Layer 1a: Message is synonym-normalised → tokenised → lemmatised → converted to BoW vector → MLP predicts intent probabilities.",
    "If max(probability) ≥ 0.65 → matching FAQ response returned immediately.",
    "Layer 1b: If MLP confidence is low → Sentence Transformer encodes query → cosine similarity computed against all pre-encoded FAQ patterns → if best score ≥ 0.50, response returned.",
    "Layer 2: If both Layer 1 attempts fail → document_ingestion.search_documents() scans active document chunks → keyword-overlap scoring → if top score ≥ 2, document excerpt returned with source attribution.",
    "Fallback: If all layers fail → query logged to unanswered_queries.json → default message returned.",
    "Flask returns JSON response → JavaScript renders it in the chat window.",
]:
    add_bullet(step)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  7. TECHNOLOGIES USED
# ═════════════════════════════════════════════════════════════════
add_main_heading(7, "TECHNOLOGIES USED")

add_table(
    ["#", "Name", "Purpose", "Why Chosen"],
    [
        ["1",  "Python 3.x",               "Core backend language",                         "Extensive AI/ML ecosystem; Flask compatibility"],
        ["2",  "Flask",                     "Web framework & routing",                       "Lightweight, Jinja2 templating, simple REST APIs"],
        ["3",  "scikit-learn (MLP)",        "Intent classification neural network",          "No GPU required; easily serialisable with Pickle"],
        ["4",  "NLTK",                      "Tokenisation & lemmatisation",                  "Industry-standard NLP preprocessing library"],
        ["5",  "Sentence Transformers",     "Semantic similarity fallback",                  "Pre-trained embeddings; handles paraphrase variations"],
        ["6",  "NumPy",                     "Bag-of-Words array operations",                 "Fast vectorised numerical computing"],
        ["7",  "pdfplumber",                "PDF text extraction",                           "Reliable textual PDF parsing; no dependency on Java"],
        ["8",  "python-docx",               "DOCX text extraction",                          "Standard library for Microsoft Word document parsing"],
        ["9",  "HTML5 + CSS3",              "Frontend UI structure & styling",               "Native browser support; no framework dependency"],
        ["10", "Vanilla JavaScript",        "AJAX chatbot message handling",                 "Lightweight; avoids heavy frontend frameworks"],
        ["11", "Jinja2",                    "HTML templating engine",                        "Built into Flask; supports dynamic page rendering"],
        ["12", "JSON",                      "Flat-file data store",                          "Portable; no database server required for deployment"],
        ["13", "Werkzeug",                  "Secure filename handling",                      "Part of Flask ecosystem; prevents path traversal"],
        ["14", "Pickle",                    "Model & vocabulary serialisation",              "Standard Python object persistence format"],
        ["15", "UUID",                      "Unique document ID generation",                 "Collision-free identifiers for document management"],
        ["16", "Regular Expressions (re)", "Text cleaning & synonym normalisation",          "Flexible pattern matching for preprocessing"],
    ],
    col_widths=[0.3, 1.6, 1.9, 2.4]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  8. MODULES DESCRIPTION
# ═════════════════════════════════════════════════════════════════
add_main_heading(8, "MODULES DESCRIPTION")

modules = [
    {
        "num": "Module 1",
        "name": "Student Chatbot Interface",
        "desc": "The public-facing chat page where students interact with the chatbot in real time. "
                "Built using HTML, CSS, and Vanilla JavaScript with AJAX for async communication.",
        "func": [
            "Renders a clean conversation window with timestamped message bubbles.",
            "Captures student input on click or Enter key press.",
            "Sends AJAX POST request to /get-response and displays the response without page reload.",
        ],
        "inputs":  "Student's natural language text (typed query).",
        "outputs": "Chatbot-generated response displayed in the chat window.",
    },
    {
        "num": "Module 2",
        "name": "Chatbot Engine (NLP Core)",
        "desc": "Central intelligence of the system. Orchestrates the two-layer hybrid response strategy "
                "via synonym normalisation, BoW vectorisation, MLP classification, sentence embeddings, "
                "and document search.",
        "func": [
            "Applies synonym normalisation (register → apply, pay → fee, etc.).",
            "Converts query to Bag-of-Words and runs MLP inference (Layer 1a).",
            "Applies Sentence Transformer cosine similarity (Layer 1b fallback).",
            "Invokes keyword document chunk search (Layer 2).",
            "Logs unresolved queries with confidence score and timestamp.",
        ],
        "inputs":  "Raw user message string.",
        "outputs": "Response string (from FAQ, document excerpt, or fallback message).",
    },
    {
        "num": "Module 3",
        "name": "NLP Model Trainer",
        "desc": "Offline script used once (or after adding new FAQs) to train the MLP classifier on the "
                "intents dataset and persist all artefacts to disk.",
        "func": [
            "Reads intents.json and extracts all question patterns and tags.",
            "Tokenises and lemmatises all patterns using clean_up_sentence().",
            "Builds a Bag-of-Words training matrix and one-hot encoded label matrix.",
            "Trains MLPClassifier(hidden_layers=(128,64), activation='relu', solver='adam').",
            "Saves model.pkl, words.pkl, classes.pkl using Pickle.",
        ],
        "inputs":  "intents.json (FAQ patterns and intent tags).",
        "outputs": "model.pkl, words.pkl, classes.pkl.",
    },
    {
        "num": "Module 4",
        "name": "NLP Preprocessor",
        "desc": "Converts raw text into standardised numerical tokens using NLTK with a graceful "
                "fallback to simple string processing if NLTK resources are unavailable.",
        "func": [
            "Tokenises sentences using nltk.word_tokenize().",
            "Lemmatises tokens using WordNetLemmatizer.",
            "Constructs binary Bag-of-Words NumPy array against the training vocabulary.",
        ],
        "inputs":  "Raw sentence string + vocabulary list (words.pkl).",
        "outputs": "Bag-of-Words NumPy array of shape [1, vocab_size].",
    },
    {
        "num": "Module 5",
        "name": "Document Ingestion & Search",
        "desc": "Handles the complete lifecycle of admin-uploaded documents: text extraction, cleaning, "
                "chunking, indexed storage, status management, deletion, and keyword search.",
        "func": [
            "Extracts text from TXT (UTF-8 decode), PDF (pdfplumber), DOCX (python-docx).",
            "Cleans text by removing excessive whitespace and non-standard characters.",
            "Splits text into ~500-character chunks at sentence boundaries.",
            "Saves metadata to admin_static_metadata.json and content to admin_static_knowledge.json.",
            "Searches active document chunks using query keyword overlap scoring.",
        ],
        "inputs":  "Uploaded file object, filename, category label.",
        "outputs": "Indexed document chunks in JSON; keyword-scored search results.",
    },
    {
        "num": "Module 6",
        "name": "Admin Authentication",
        "desc": "Session-based login system that protects all admin routes using a decorator pattern. "
                "All admin pages require an active authenticated session to access.",
        "func": [
            "Validates username and password against hardcoded admin credentials.",
            "Sets 'admin_logged_in' flag in Flask session on success.",
            "Applies @login_required decorator to redirect unauthenticated requests.",
            "Clears session on logout.",
        ],
        "inputs":  "Username and password from the admin login form.",
        "outputs": "Authenticated session or redirect with error flash message.",
    },
    {
        "num": "Module 7",
        "name": "FAQ Management",
        "desc": "Allows administrators to add new FAQ intent entries and delete existing ones through a "
                "web interface, directly modifying the intents.json knowledge base.",
        "func": [
            "Presents a paginated list of all existing intents with pattern previews.",
            "Parses multi-line admin input into patterns and responses lists.",
            "Validates for duplicate tags before insertion.",
            "Writes updated intent list back to intents.json.",
        ],
        "inputs":  "Intent tag (string), patterns (newline-separated), responses (newline-separated).",
        "outputs": "Updated intents.json; success or error flash message.",
    },
    {
        "num": "Module 8",
        "name": "Analytics Engine",
        "desc": "Aggregates and computes performance metrics from stored query logs, the FAQ knowledge "
                "base, and uploaded document metadata to provide actionable insights.",
        "func": [
            "Counts total queries, FAQ-answered, document-answered, and unanswered.",
            "Calculates success_rate = (faq + doc answered) / total * 100.",
            "Identifies and ranks top-10 most frequently unanswered queries.",
            "Counts total FAQs, total uploaded documents, active documents, and total chunks.",
        ],
        "inputs":  "unanswered_queries.json, intents.json, admin_static_metadata.json.",
        "outputs": "Analytics dictionary rendered on the admin analytics page.",
    },
]

for mod in modules:
    add_sub_heading(f"{mod['num']}: {mod['name']}")
    add_paragraph(mod["desc"])
    add_paragraph("Functionality:", bold=True, size=10.5, color=NAVY, space_after=2)
    for f in mod["func"]:
        add_bullet(f)
    p = doc.add_paragraph()
    r1 = p.add_run("Input:  "); r1.bold = True; r1.font.color.rgb = TEAL; r1.font.size = Pt(10.5)
    r2 = p.add_run(mod["inputs"]); r2.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    r3 = p2.add_run("Output: "); r3.bold = True; r3.font.color.rgb = TEAL; r3.font.size = Pt(10.5)
    r4 = p2.add_run(mod["outputs"]); r4.font.size = Pt(10.5)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  9. WORKING OF THE SYSTEM
# ═════════════════════════════════════════════════════════════════
add_main_heading(9, "WORKING OF THE SYSTEM")

add_sub_heading("Step-by-Step Working Process")
steps = [
    ("Step 1 – Server Startup", "Admin runs python app.py. Flask initialises, chatbot_engine.py loads intents.json, deserialises model.pkl / words.pkl / classes.pkl, and pre-computes Sentence Transformer embeddings for every FAQ pattern."),
    ("Step 2 – Student Visits Homepage", "Student navigates to http://localhost:5000/. The main.html college homepage is served with navigation links to the Student Helpdesk and Admin Login."),
    ("Step 3 – Chatbot Page Load", "Student clicks 'Student Helpdesk' → /helpdesk → chatbot.html is rendered displaying the chat interface."),
    ("Step 4 – Query Submission", "Student types a question and presses Send/Enter. JavaScript captures the input, constructs a JSON payload, and sends an AJAX POST to /get-response without reloading the page."),
    ("Step 5 – Layer 1a: MLP Classification", "The message is synonym-normalised, tokenised by NLTK, lemmatised, and converted to a Bag-of-Words vector. The MLP classifier returns probabilities for all 40+ intent classes. If max(probability) ≥ 0.65, the matched FAQ response is returned."),
    ("Step 6 – Layer 1b: Semantic Embedding Fallback", "If MLP confidence is below the threshold, the query is encoded into a 384-dimensional vector and compared against all pre-computed pattern embeddings using cosine similarity. If the highest score ≥ 0.50, the response for the matched intent is returned."),
    ("Step 7 – Layer 2: Document Search", "If both Layer 1 attempts fail, search_documents() scans all active document chunks. Each chunk is scored by the number of query keywords it contains. The highest-scoring chunk is returned with source attribution if its score ≥ 2."),
    ("Step 8 – Fallback & Logging", "If no layer resolves the query, a default fallback message is returned and the query is logged to data/unanswered_queries.json with timestamp, confidence score, and source='none'."),
    ("Step 9 – Admin Review & Improvement", "Admin logs into /admin/login → views Analytics page to identify frequent failures → adds new FAQ intents or uploads new documents → optionally re-trains the model by running train_model.py."),
]
for title, body in steps:
    add_paragraph(title, bold=True, size=11, color=NAVY, space_before=6, space_after=2)
    add_paragraph(body, space_after=4)

add_sub_heading("User Flow Summary")
add_code_block(
"""Student opens /helpdesk
    │
    ▼  Types query
JavaScript  →  AJAX POST  →  Flask /get-response
                                    │
                     ┌──────────────▼──────────────┐
                     │     LAYER 1a – MLP           │  confidence ≥ 0.65 ?
                     │     Bag-of-Words + model     │──── YES ──→ FAQ Response
                     └──────────────┬──────────────┘
                                    │ NO
                     ┌──────────────▼──────────────┐
                     │  LAYER 1b – Sentence Embed  │  similarity ≥ 0.50 ?
                     │  Cosine Similarity Search   │──── YES ──→ Embedding Response
                     └──────────────┬──────────────┘
                                    │ NO
                     ┌──────────────▼──────────────┐
                     │   LAYER 2 – Document Search │  keyword score ≥ 2 ?
                     │   Chunk Keyword Matching    │──── YES ──→ Document Response
                     └──────────────┬──────────────┘
                                    │ NO
                              Log query to JSON
                              Return fallback message"""
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  10. ALGORITHMS / LOGIC USED
# ═════════════════════════════════════════════════════════════════
add_main_heading(10, "ALGORITHMS / LOGIC USED")

algos = [
    ("Algorithm 1: Bag-of-Words (BoW) Vectorisation",
     "Converts text to a fixed-length binary vector for MLP input.",
     ["Tokenise input using NLTK word_tokenize().",
      "Lemmatise all tokens using WordNetLemmatizer.",
      "Initialise a zero vector of length = |vocabulary|.",
      "Set index[i] = 1 if vocabulary word i appears in the lemmatised tokens.",
      "Return the binary NumPy array."]),
    ("Algorithm 2: MLP Intent Classification (scikit-learn)",
     "Classifies student intent using a multi-layer perceptron with softmax output.",
     ["Architecture: Input Layer → Dense(128, ReLU) → Dense(64, ReLU) → Output(n_classes, Softmax).",
      "Solver: Adam. Max Iterations: 300. Random State: 1.",
      "Acceptance Threshold: max(probability) ≥ 0.65.",
      "If accepted, the top predicted class label maps to an intent in intents.json.",
      "A random response is selected from the matched intent's 'responses' list."]),
    ("Algorithm 3: Sentence Embedding Cosine Similarity",
     "Handles semantically similar queries not covered by BoW vocabulary.",
     ["At startup, encode all FAQ patterns using all-MiniLM-L6-v2 → 384-dim vectors stored in a matrix.",
      "At query time, encode the user query → compute cosine_sim(query, matrix).",
      "Find argmax(cosine_scores) → best matching pattern index.",
      "If best_score ≥ 0.50, return a response from the matched intent.",
      "Formula: similarity(A,B) = (A · B) / (‖A‖ × ‖B‖)"]),
    ("Algorithm 4: Document Keyword Overlap Search",
     "Retrieves the most relevant document chunk using word-level overlap.",
     ["Tokenise query into a set of unique words.",
      "For each chunk of each active document, tokenise the chunk into a word set.",
      "Score = |query_words ∩ chunk_words| (set intersection size).",
      "Sort all scored chunks in descending order.",
      "Return top-K chunks; accept if score ≥ DOCUMENT_RELEVANCE_THRESHOLD (= 2)."]),
    ("Algorithm 5: Synonym Normalisation",
     "Pre-processing step to standardise vocabulary before ML inference.",
     ["Define a synonym dictionary: {register→apply, pay→fee, semester→term, course→program}.",
      "For each synonym pair, apply whole-word regex replacement on lowercased input.",
      r"Pattern: re.sub(r'\b' + word + r'\b', replacement, sentence).",
      "Whole-word matching prevents accidental replacements (e.g., 'layer' → not affected)."]),
    ("Algorithm 6: Text Chunking",
     "Splits large document text into retrievable 500-character segments.",
     ["Split text by sentence boundaries using regex: r'(?<=[.!?])\\s+'.",
      "Accumulate sentences into a current_chunk buffer.",
      "When len(current_chunk) + len(next_sentence) ≥ 500, save chunk and start a new one.",
      "Append the final buffer as the last chunk."]),
]

for name, purpose, steps_ in algos:
    add_sub_heading(name)
    add_paragraph(f"Purpose: {purpose}", italic=True)
    for s in steps_:
        add_bullet(s)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  11. DATABASE DESIGN
# ═════════════════════════════════════════════════════════════════
add_main_heading(11, "DATABASE DESIGN")
add_paragraph(
    "The project uses JSON flat-files as a lightweight data store, eliminating the need for an "
    "external database server. All files reside within the project directory."
)

schemas = [
    ("File 1: intents.json — FAQ Knowledge Base",
     ["Field", "Type", "Description"],
     [["intents", "Array", "Top-level container holding all intent objects"],
      ["tag", "String", "Unique intent identifier (e.g., 'admissions', 'library')"],
      ["patterns", "Array[String]", "Sample user questions used to train the classifier"],
      ["responses", "Array[String]", "Possible chatbot responses; one is selected randomly"]]),
    ("File 2: data/admin_static_metadata.json — Document Index",
     ["Field", "Type", "Description"],
     [["documents", "Array", "List of all uploaded document metadata objects"],
      ["doc_id", "String (UUID)", "Unique collision-free document identifier"],
      ["title", "String", "Original uploaded filename"],
      ["category", "String", "Admin-assigned category (Fees, Exams, Events, etc.)"],
      ["uploaded_at", "String (datetime)", "ISO-format upload timestamp"],
      ["status", "String", "'active' or 'inactive' (controls search inclusion)"],
      ["total_chunks", "Integer", "Number of text chunks extracted from the document"]]),
    ("File 3: data/admin_static_knowledge.json — Document Content",
     ["Field", "Type", "Description"],
     [["doc_id", "String (UUID)", "Foreign key linking to the metadata file"],
      ["extracted_text_content", "String", "Full cleaned text extracted from the document"],
      ["chunks", "Array[String]", "Pre-split sentence-boundary chunks for fast retrieval"]]),
    ("File 4: data/unanswered_queries.json — Query Audit Log",
     ["Field", "Type", "Description"],
     [["query", "String", "Original user query text"],
      ["confidence", "Float", "Max probability from MLP output (0.0 – 1.0)"],
      ["source", "String", "'document' (partial success) or 'none' (complete failure)"],
      ["status", "String", "'unanswered' or 'resolved' (admin marked)"],
      ["timestamp", "String (datetime)", "When the query was logged"]]),
]

for title, headers, rows in schemas:
    add_sub_heading(title, size=11)
    add_table(headers, rows, col_widths=[2.0, 1.5, 3.1])

add_sub_heading("Entity Relationships")
add_code_block(
"""intents.json  ──[trained on patterns]──▶  model.pkl / words.pkl / classes.pkl
admin_static_metadata.json  ──[doc_id]──▶  admin_static_knowledge.json
unanswered_queries.json  ──[reviewed by admin]──▶  admin adds to intents.json / uploads docs"""
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  12. UML DIAGRAMS (TEXT EXPLANATIONS)
# ═════════════════════════════════════════════════════════════════
add_main_heading(12, "UML DIAGRAMS (EXPLAINED IN TEXT)")

# Use Case
add_sub_heading("12.1  Use Case Diagram")
add_paragraph("Actors: Student, Admin", bold=True)
add_paragraph("Student Use Cases:")
for uc in ["View College Homepage", "Open Student Helpdesk", "Send Query to Chatbot", "Receive Chatbot Response"]:
    add_bullet(uc)
add_paragraph("Admin Use Cases:")
for uc in ["Login to Admin Panel", "View Dashboard", "Add FAQ Intent", "Delete FAQ Intent",
           "Upload Document", "Toggle Document Active/Inactive", "Delete Document",
           "View Analytics & Knowledge Gaps", "Logout"]:
    add_bullet(uc)
add_paragraph("Relationships:", bold=True)
add_bullet('Student "Send Query" extends "Receive Response"')
add_bullet('All Admin use cases include "Login to Admin Panel" (dependency)')
add_bullet('"Add FAQ" and "Upload Document" both include "View Knowledge Page"')

add_sub_heading("12.2  Class Diagram")
classes_desc = [
    ("ChatbotEngine", 
     "words, classes, model, intents, embedding_model, intent_embeddings, CONFIDENCE_THRESHOLD",
     "normalize_synonyms(), predict_class(), get_faq_response(), get_embedding_response(), search_admin_static_knowledge(), chatbot_response()"),
    ("DocumentIngestion",
     "(stateless module — no instance attributes)",
     "extract_text_from_txt(), extract_text_from_pdf(), extract_text_from_docx(), clean_text(), chunk_text(), process_document(), save_document_data(), load_knowledge_base(), search_documents(), update_document_status(), delete_document()"),
    ("NLPPreprocessor",
     "lemmatizer, USE_NLTK",
     "clean_up_sentence(), bag_of_words()"),
    ("FlaskApp",
     "app, ADMIN_USERNAME, ADMIN_PASSWORD, DATA_DIR, UPLOAD_FOLDER",
     "home(), helpdesk(), get_bot_response(), admin_login(), admin_logout(), admin_dashboard(), admin_analytics(), admin_knowledge(), admin_add_faq(), admin_delete_faq(), upload_document(), compute_analytics()"),
]
add_table(
    ["Class", "Attributes", "Methods"],
    [[c, a, m] for c, a, m in classes_desc],
    col_widths=[1.5, 2.1, 3.0]
)
add_paragraph("Relationships:", bold=True)
add_bullet("FlaskApp USES ChatbotEngine (calls chatbot_response())")
add_bullet("ChatbotEngine USES NLPPreprocessor (calls bag_of_words(), clean_up_sentence())")
add_bullet("ChatbotEngine USES DocumentIngestion (calls search_documents())")
add_bullet("FlaskApp USES DocumentIngestion (calls save_document_data(), load_knowledge_base())")

add_sub_heading("12.3  Sequence Diagram: Student Query Flow")
add_code_block(
"""Student     Browser       Flask         ChatbotEngine    NLPPreprocessor  MLPModel  SentenceTF  DocIngestion
   │            │             │                  │                │             │           │            │
   │─ type ────▶│             │                  │                │             │           │            │
   │            │─ AJAX POST ─▶/get-response     │                │             │           │            │
   │            │             │─ chatbot_resp() ─▶               │             │           │            │
   │            │             │                  │─ bag_of_words()▶             │           │            │
   │            │             │                  │                │─ return ────▶            │            │
   │            │             │                  │─ predict_proba()─────────────▶            │            │
   │            │             │                  │                │            ─▶ probs      │            │
   │            │             │     [conf≥0.65]  ◀─ FAQ response ─┤            │            │            │
   │            │             │     [conf<0.65]  │─ encode() ──────────────────────────────▶│            │
   │            │             │                  │◀─ cosine score ─────────────────────────┤            │
   │            │             │     [sim≥0.50]   ◀─ embed response┤            │            │            │
   │            │             │     [sim<0.50]   │─ search_documents()──────────────────────────────────▶│
   │            │             │                  │◀─ scored chunks ─────────────────────────────────────┤
   │            │             │     [score≥2]    ◀─ doc response  ─┤            │            │            │
   │            │             │     [score<2]    │─ log_query() / return fallback             │            │
   │            │◀─ JSON resp ─┤                  │                │             │            │            │
   │◀─ render  ─┤             │                  │                │             │            │            │"""
)

add_sub_heading("12.4  Activity Diagram: Admin Document Upload")
add_code_block(
"""[Start]
    ↓
Admin selects file + category and submits form
    ↓
Is file present?
  ├── No  →  Flash error  →  [End]
  └── Yes →  Read file in-memory (TXT / PDF / DOCX)
                ↓
         Is content extractable?
           ├── No (scanned PDF / empty)  →  Flash error  →  [End]
           └── Yes  →  Clean text (regex whitespace normalisation)
                            ↓
                       Chunk text (~500 char / chunk, sentence boundaries)
                            ↓
                       Generate UUID + upload timestamp
                            ↓
                       Build metadata entry (doc_id, title, category, status, total_chunks)
                            ↓
                       Build content entry (doc_id, extracted_text, chunks)
                            ↓
                       Save to admin_static_metadata.json
                            ↓
                       Save to admin_static_knowledge.json
                            ↓
                       Flash success message
                            ↓
                        [End]"""
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  13. USER INTERFACE DESCRIPTION
# ═════════════════════════════════════════════════════════════════
add_main_heading(13, "USER INTERFACE DESCRIPTION")

pages = [
    ("Page 1: College Homepage  (/)",
     "main.html — Public landing page for Pallavi Engineering College.",
     ["Navigation Bar: College logo + links (Home, Student Helpdesk, Admin Login).",
      "Hero Section: Full-width banner with college tagline and mission statement.",
      "About Section: Brief description of the institution, affiliations, and accreditations.",
      "Stats Grid: Cards for Affiliation, Accreditations, and Location.",
      "Campus Facilities Gallery: Cards showing Auditorium, Digital Library, and Robotics Lab.",
      "Footer: Address, copyright, and contact information."]),
    ("Page 2: Student Helpdesk  (/helpdesk)",
     "chatbot.html — The student-facing conversational chatbot interface.",
     ["Chat message window with scrollable conversation history.",
      "Student messages appear on the right (primary colour bubble).",
      "Bot messages appear on the left (neutral grey bubble).",
      "Text input field and Send button at the bottom.",
      "JavaScript sends AJAX request and renders response without page reload."]),
    ("Page 3: Admin Login  (/admin/login)",
     "admin/login.html — Secure access gate for admin panel.",
     ["Username and password input fields.",
      "Login submit button.",
      "Flash message area for invalid credentials error."]),
    ("Page 4: Admin Dashboard  (/admin/dashboard)",
     "admin/dashboard.html — Admin home screen with navigation cards.",
     ["Quick-access cards: FAQ Management, Document Upload, Analytics.",
      "Logout button.",
      "Summary statistics (optional): total FAQs, total documents."]),
    ("Page 5: Knowledge Management  (/admin/knowledge)",
     "admin/knowledge.html — Combined FAQ and document management page.",
     ["FAQ Tab: Table of all intent tags with pattern previews; Add FAQ form (tag, patterns, responses).",
      "Document Tab: Upload form with file chooser and category dropdown.",
      "Document table: filename, category, upload date, status, chunk count, Toggle/Delete actions."]),
    ("Page 6: Analytics Dashboard  (/admin/analytics)",
     "admin/analytics.html — Performance insights and knowledge gap analysis.",
     ["Stat Banner (4 cards): Total Queries, Success Rate, FAQ Answers, Unresolved.",
      "Knowledge Gaps Table: Top 10 frequent unanswered queries with frequency count.",
      "Improvement Insights: Conditional alerts when success rate < 80% or gaps exist.",
      "Source Distribution: FAQ-driven vs. Document-driven vs. Fallback counts."]),
]

for page_title, desc, elements in pages:
    add_sub_heading(page_title)
    add_paragraph(desc, italic=True)
    for e in elements:
        add_bullet(e)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  14. FEATURES
# ═════════════════════════════════════════════════════════════════
add_main_heading(14, "FEATURES OF THE PROJECT")

features = [
    ("Two-Layer Hybrid Chatbot",        "MLP classifier + document fallback ensures maximum query coverage."),
    ("Semantic Query Matching",          "Sentence Transformer cosine similarity handles paraphrased or novel questions."),
    ("Synonym Normalisation",           "Pre-processing maps common synonyms before ML inference for consistency."),
    ("Natural Language Understanding",  "NLTK tokenisation and WordNet lemmatisation for vocabulary normalisation."),
    ("Multi-Format Document Upload",    "Supports PDF, DOCX, and TXT file ingestion with automatic text extraction."),
    ("Dynamic Document Management",     "Admins can activate or deactivate documents without deleting them."),
    ("FAQ CRUD via Web UI",             "Admins add and delete FAQ intents through a browser form — no coding needed."),
    ("Unanswered Query Logging",        "All failed queries are stored with timestamp and confidence score."),
    ("Analytics Dashboard",             "Real-time success rate, source distribution, and knowledge gap analysis."),
    ("Session-Based Admin Auth",        "Secure login/logout with session flag and login_required decorator."),
    ("Document Category Tagging",       "Documents are tagged by topic (Fees, Exams, Events) for better retrieval."),
    ("Smart Text Chunking",             "Sentence-boundary splitting ensures chunks are coherent and searchable."),
    ("Scanned PDF Detection",           "System warns admin when a PDF has no extractable text (image-based)."),
    ("File Size Validation",            "5 MB maximum upload size enforced to prevent server overload."),
    ("Real-Time Chat Interface",        "AJAX-based communication shows responses instantly without page refresh."),
]

add_table(
    ["#", "Feature", "Description"],
    [[str(i+1), f, d] for i, (f, d) in enumerate(features)],
    col_widths=[0.3, 2.0, 4.3]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  15. ADVANTAGES
# ═════════════════════════════════════════════════════════════════
add_main_heading(15, "ADVANTAGES")

add_sub_heading("Technical Advantages")
for a in [
    "No GPU Required: MLPClassifier runs entirely on CPU — deployable on any standard server.",
    "Modular Architecture: Each component (NLP, document search, analytics) is independently testable and replaceable.",
    "No External Database Server: JSON file store makes setup, backup, and migration trivial.",
    "Graceful Degradation: Three-layer fallback prevents hard failures — system always returns a response.",
    "Semantic Robustness: Sentence Transformers handle vocabulary variation not covered by the BoW model.",
    "Lightweight Footprint: Total application size is under 2 MB (excluding Pickle model and Python packages).",
]:
    add_bullet(a)

add_sub_heading("Practical Advantages")
for a in [
    "Zero Technical Skill Required from Admins: FAQ/document updates done entirely through browser forms.",
    "Instant Deployment: python app.py — no cloud infrastructure or database provisioning needed.",
    "Scalable Knowledge Base: New FAQs and documents can be added live without redeployment or downtime.",
    "Transparent Performance: Analytics dashboard shows exactly where the system succeeds and fails.",
    "Consistent Student Experience: Same answer quality 24/7, regardless of staff availability.",
    "Data-Driven Improvement: Unanswered query logs guide admins toward the highest-impact improvements.",
]:
    add_bullet(a)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  16. LIMITATIONS
# ═════════════════════════════════════════════════════════════════
add_main_heading(16, "LIMITATIONS")

add_table(
    ["#", "Limitation", "Impact"],
    [
        ["1", "Static Model — new FAQs require manual retraining", "Admin must run train_model.py after adding new intents"],
        ["2", "Scanned PDFs not supported (no OCR)", "Image-based institutional documents cannot be ingested"],
        ["3", "Keyword-only Layer 2 search", "Document search may miss relevant chunks using different vocabulary"],
        ["4", "Hardcoded admin credentials in app.py", "Not suitable for production without a proper auth system"],
        ["5", "Single admin account", "No multi-user or role-based admin access control"],
        ["6", "No persistent conversation history", "Chat session lost on browser refresh"],
        ["7", "JSON file store without indexing", "Performance may degrade with very large query log files"],
        ["8", "English language only", "No support for Telugu, Hindi, or other regional languages"],
        ["9", "5 MB maximum upload size", "Large institutional documents must be split before upload"],
    ],
    col_widths=[0.3, 2.8, 3.5]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  17. FUTURE ENHANCEMENTS
# ═════════════════════════════════════════════════════════════════
add_main_heading(17, "FUTURE ENHANCEMENTS")

add_sub_heading("Short-Term Improvements")
for e in [
    "Automatic Model Retraining: Trigger train_model.py automatically when new FAQs are saved via the admin panel.",
    "Password Hashing: Store admin credentials using bcrypt/Argon2 in an SQLite user table.",
    "Multi-Admin Support: Implement role-based access control with separate admin accounts.",
    "OCR Integration: Use Tesseract-OCR or EasyOCR to extract text from scanned image PDFs.",
    "Semantic Layer 2 Search: Replace keyword overlap with vector embeddings for document retrieval.",
]:
    add_bullet(e)

add_sub_heading("Medium-Term Improvements")
for e in [
    "Persistent Chat History: Store per-session conversations in SQLite for audit and replay.",
    "Multi-Language Support: Add Telugu and Hindi intent sets for regional accessibility.",
    "WhatsApp / Telegram Integration: Deploy the chatbot on popular messaging platforms via webhooks.",
    "Email / SMS Alerts: Notify administrators when the unanswered query count exceeds a threshold.",
    "REST API Layer: Expose the chatbot as a REST API for integration with college mobile applications.",
]:
    add_bullet(e)

add_sub_heading("Long-Term Scalability")
for e in [
    "Vector Database (ChromaDB / Pinecone): Replace JSON chunk search with a proper vector store for thousands of document segments.",
    "LLM Integration: Use a hosted Large Language Model (Gemini, GPT-4) for free-form generation when FAQ and document search both fail.",
    "Real-Time Fine-Tuning: Automatically incorporate admin-resolved queries into the training dataset.",
    "Multi-College Deployment: Containerise with Docker and deploy across multiple institutions on a shared cloud platform.",
    "Voice Interface: Add speech-to-text input for accessibility and kiosk-style deployment.",
]:
    add_bullet(e)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  18. TESTING
# ═════════════════════════════════════════════════════════════════
add_main_heading(18, "TESTING")

add_sub_heading("Types of Testing Used")
add_table(
    ["Type", "Description", "Tool / Method"],
    [
        ["Unit Testing",        "Individual functions tested in isolation",                                      "test_suite.py (custom script)"],
        ["Integration Testing", "Flask routes tested with simulated HTTP requests",                            "Flask test client"],
        ["Functional Testing",  "End-to-end chatbot response flow validated",                                  "Manual + scripted inputs"],
        ["Manual Testing",      "UI interaction validated in the browser",                                      "Chrome DevTools"],
        ["Negative Testing",    "Invalid inputs, empty messages, wrong credentials, unsupported file types",   "Manual test cases"],
    ],
    col_widths=[1.5, 2.8, 2.3]
)

add_sub_heading("Test Cases")
add_table(
    ["#", "Test Case", "Input", "Expected Output", "Result"],
    [
        ["1",  "Greeting intent",             '"Hello"',                             "Welcome message from FAQ",                               "Pass"],
        ["2",  "Admission query",             '"How do I apply for admission?"',     "Admission process details from FAQ",                     "Pass"],
        ["3",  "Library timing",              '"When does the library close?"',      "8:00 AM – 6:00 PM working days",                        "Pass"],
        ["4",  "Synonym normalisation",       '"How do I register?"',               "Admission response (register → apply)",                  "Pass"],
        ["5",  "Low confidence fallback",     '"blah blah blah"',                   "Default fallback message + query logged",                 "Pass"],
        ["6",  "Document search hit",         '"fee deadline" (doc uploaded)',       "Based on fee_doc.pdf: [excerpt]",                        "Pass"],
        ["7",  "Admin login — valid",         "Correct username & password",        "Redirect to /admin/dashboard",                           "Pass"],
        ["8",  "Admin login — invalid",       "Wrong password",                     '"Invalid credentials" flash error',                      "Pass"],
        ["9",  "Add FAQ intent",              "Valid tag + patterns + responses",   "Intent added to intents.json; success flash",            "Pass"],
        ["10", "Delete FAQ intent",           "Existing tag selected",              "Intent removed; success flash",                          "Pass"],
        ["11", "Upload PDF",                  "Valid text-based PDF",               "'Success! Processed and indexed.' flash",                 "Pass"],
        ["12", "Upload scanned PDF",          "Image-only PDF",                     "Error: 'No extractable text' flash",                     "Pass"],
        ["13", "Upload oversized file",       "File > 5 MB",                        "413 Request Entity Too Large error",                     "Pass"],
        ["14", "Analytics computation",       "10 entries in query log",            "Correct counts and success rate displayed",              "Pass"],
        ["15", "Protected route without auth","GET /admin/dashboard (no session)",  "Redirect to /admin/login",                               "Pass"],
    ],
    col_widths=[0.3, 1.8, 1.9, 2.0, 0.6]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  19. RESULTS
# ═════════════════════════════════════════════════════════════════
add_main_heading(19, "RESULTS")

add_sub_heading("Output for Students")
for r in [
    "Instant, natural language responses to 40+ college-specific FAQ categories including admissions, exams, fees, library, placements, hostel, and contacts.",
    "Context-sensitive answers extracted from admin-uploaded institutional documents with clear source attribution.",
    "A graceful fallback message with a promise of admin review when knowledge is unavailable.",
    "Responses delivered in under 500 ms for FAQ lookups; document search typically completes within 1–2 seconds.",
]:
    add_bullet(r)

add_sub_heading("Output for Administrators")
for r in [
    "Real-time analytics dashboard showing total queries, success rate percentage, and source breakdown (FAQ / Document / Fallback).",
    "Ranked list of top-10 most frequently unanswered queries, pinpointing exact knowledge gaps.",
    "Conditional improvement alerts when success rate drops below 80%.",
    "Full audit log of unanswered queries with timestamps and confidence scores for review.",
]:
    add_bullet(r)

add_sub_heading("System Performance Characteristics")
add_table(
    ["Metric", "Value / Description"],
    [
        ["MLP Confidence Threshold",       "0.65 — ensures high-precision FAQ responses"],
        ["Semantic Similarity Threshold",  "0.50 cosine score — handles paraphrase variation"],
        ["Document Relevance Threshold",   "≥ 2 keyword overlaps — prevents irrelevant document snippets"],
        ["FAQ Intents Supported",          "40+ college-specific intent categories in intents.json"],
        ["Supported Upload Formats",       "PDF (text-based), DOCX, TXT"],
        ["Maximum Upload Size",            "5 MB per document"],
        ["Sentence Embedding Model",       "all-MiniLM-L6-v2 (384-dimensional vectors)"],
        ["Neural Network Architecture",    "MLP: 128 → 64 → N_classes (ReLU, Adam, Softmax)"],
    ],
    col_widths=[2.5, 4.1]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  20. CONCLUSION
# ═════════════════════════════════════════════════════════════════
add_main_heading(20, "CONCLUSION")

add_paragraph(
    "The AI-Powered College Helpdesk Chatbot successfully addresses a critical institutional challenge: "
    "providing students with reliable, instant, and accurate information regardless of time or staff "
    "availability. By integrating a trained MLP neural network, semantic sentence embeddings, and an "
    "admin-driven document retrieval system into a single unified chatbot, the project demonstrates how "
    "modern AI techniques can be applied practically and cost-effectively in an educational setting."
)
add_paragraph(
    "The two-layer hybrid architecture ensures that queries are answered with maximum coverage — the "
    "neural network handles high-frequency, well-defined questions with strong confidence, while the "
    "sentence embedding fallback and document search layer capture edge cases and dynamic institutional "
    "content that cannot be pre-programmed. The admin panel completes the system by creating a "
    "sustainable feedback loop: every unanswered query becomes a data point that guides targeted "
    "improvement through new FAQs or document uploads."
)
add_paragraph(
    "Built without proprietary infrastructure — using only Python, Flask, scikit-learn, NLTK, and "
    "JSON file storage — the system is portable, lightweight, and immediately deployable in any "
    "college environment. It lays a strong foundation for future enhancements including LLM integration, "
    "vector database search, voice interfaces, and multi-college deployment, ensuring its relevance "
    "well into the future of AI-driven education management."
)
add_paragraph(
    "In summary, this project is not merely a chatbot; it is a self-improving institutional knowledge "
    "management system that transforms student-facing information delivery from reactive and manual to "
    "proactive, intelligent, and data-driven.",
    bold=False, italic=True
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
#  21. REFERENCES
# ═════════════════════════════════════════════════════════════════
add_main_heading(21, "REFERENCES")

add_sub_heading("Libraries and Frameworks")
refs_libs = [
    "[1] Flask Web Framework. Pallets Projects. https://flask.palletsprojects.com/",
    "[2] scikit-learn MLPClassifier. Pedregosa et al. https://scikit-learn.org/stable/modules/generated/sklearn.neural_network.MLPClassifier.html",
    "[3] NLTK Natural Language Toolkit. Bird, S., Klein, E., & Loper, E. (2009). Natural Language Processing with Python. O'Reilly Media. https://www.nltk.org/",
    "[4] Sentence Transformers (all-MiniLM-L6-v2). https://www.sbert.net/",
    "[5] pdfplumber PDF Extraction Library. https://github.com/jsvine/pdfplumber",
    "[6] python-docx DOCX Parser. https://python-docx.readthedocs.io/",
    "[7] NumPy Scientific Computing. https://numpy.org/",
    "[8] Werkzeug WSGI Utility Library. https://werkzeug.palletsprojects.com/",
    "[9] Jinja2 Templating Engine. https://jinja.palletsprojects.com/",
]
for r in refs_libs:
    add_paragraph(r, size=10)

add_sub_heading("Academic and Research References")
refs_academic = [
    "[10] Rumelhart, D. E., Hinton, G. E., & Williams, R. J. (1986). Learning representations by back-propagating errors. Nature, 323, 533–536.",
    "[11] Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. Proceedings of EMNLP 2019. https://arxiv.org/abs/1908.10084",
    "[12] Miller, G. A. (1995). WordNet: A lexical database for English. Communications of the ACM, 38(11), 39–41.",
    "[13] Manning, C., & Schütze, H. (1999). Foundations of Statistical Natural Language Processing. MIT Press.",
    "[14] Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL 2019. https://arxiv.org/abs/1810.04805",
]
for r in refs_academic:
    add_paragraph(r, size=10)

add_sub_heading("Institutional Resources")
refs_inst = [
    "[15] Pallavi Engineering College Official Website. https://www.pallaviengineeringcollege.ac.in",
    "[16] JNTUH Academic Regulations. Jawaharlal Nehru Technological University, Hyderabad. https://jntuh.ac.in",
    "[17] AICTE Regulations for Technical Institutions. All India Council for Technical Education. https://www.aicte-india.org",
]
for r in refs_inst:
    add_paragraph(r, size=10)

# ── SAVE ──────────────────────────────────────────────────────────
output_path = r"d:\Projects\Major Project\College_Chatbot_Project_Documentation.docx"
doc.save(output_path)
print(f"Document saved successfully to:\n{output_path}")
